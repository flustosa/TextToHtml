import tkinter as tk
from tkinter import filedialog, messagebox, font, ttk
from tkinter.font import Font
import docx

def convert_to_html_text(text):
    lines = text.split("\n")
    result = ""

    for line in lines:
        if line.strip():
            result += f"<p>{line}</p>"
        else:
            result += "<p>&nbsp;</p>"

    return result

def convert_to_html_docx(docx_path):
    doc = docx.Document(docx_path)
    html_text = ""

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            html_text += f"<p>{paragraph.text}</p>"
        else:
            html_text += "<p>&nbsp;</p>"

    return html_text

def select_and_convert():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])

    if file_path:
        html_text = convert_to_html_docx(file_path)
        output_file = file_path.replace(".docx", "_converted.txt")

        with open(output_file, "w", encoding="utf-8") as file:
            file.write(html_text)

        messagebox.showinfo("Conversión completada", f"El resultado se ha guardado en '{output_file}'")

def convert_text():
    text = input_text.get("1.0", tk.END)
    if text.strip():
        html_text = convert_to_html_text(text)
        output_text.delete("1.0", tk.END)
        output_text.insert(tk.END, html_text)
    else:
        messagebox.showerror("Error", "Por favor, ingrese texto válido.")



def alignment_to_css(alignment):
    if alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    elif alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    elif alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    else:
        return "left"



def convert_to_html_docx(docx_path):
    doc = docx.Document(docx_path)
    html_text = ""

    for paragraph in doc.paragraphs:
        paragraph_html = ""

        for run in paragraph.runs:
            run_text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if run.bold and run.italic and run.underline:
                run_html = f"<strong><em><u>{run_text}</u></em></strong>"
            elif run.bold and run.italic:
                run_html = f"<strong><em>{run_text}</em></strong>"
            elif run.bold and run.underline:
                run_html = f"<strong><u>{run_text}</u></strong>"
            elif run.italic and run.underline:
                run_html = f"<em><u>{run_text}</u></em>"
            elif run.bold:
                run_html = f"<strong>{run_text}</strong>"
            elif run.italic:
                run_html = f"<em>{run_text}</em>"
            elif run.underline:
                run_html = f"<u>{run_text}</u>"
            else:
                run_html = run_text

            paragraph_html += run_html

        if paragraph_html.strip():
            html_text += f"<p>{paragraph_html}</p>"
        else:
            html_text += "<p>&nbsp;</p>"

    return html_text

def apply_style(style):
    selected_text = input_text.selection_get()
    start_index = input_text.index(tk.SEL_FIRST)
    end_index = input_text.index(tk.SEL_LAST)
    
    if not selected_text:
        return

    current_font = font.nametofont(input_text.tag_cget(start_index, "font"))

    if style == "bold":
        if current_font.actual()["weight"] == "normal":
            new_weight = "bold"
        else:
            new_weight = "normal"
        current_font.configure(weight=new_weight)

    if style == "italic":
        if current_font.actual()["slant"] == "roman":
            new_slant = "italic"
        else:
            new_slant = "roman"
        current_font.configure(slant=new_slant)

    input_text.tag_add("style", start_index, end_index)
    input_text.tag_configure("style", font=current_font)



def change_alignment():
    selected_alignment = align_var.get()

    if input_text.tag_ranges(tk.SEL):
        input_text.tag_configure(selected_alignment, justify=selected_alignment)
        input_text.tag_add(selected_alignment, input_text.index(tk.SEL_FIRST), input_text.index(tk.SEL_LAST))

app = tk.Tk()
app.title("Convertidor de Word/Texto a HTML")

# Botón para seleccionar y convertir un archivo de Word
select_button = tk.Button(app, text="Seleccionar archivo", command=select_and_convert)
select_button.pack(pady=5)

# Cuadro de texto para ingresar texto
input_label = tk.Label(app, text="Ingrese texto:")
input_label.pack(pady=5)
input_text = tk.Text(app, wrap=tk.WORD, height=10)
input_text.pack(padx=10, pady=5)

# Configurar tags para negrita, cursiva y subrayado
input_text_font = font.nametofont(input_text["font"])
input_text.tag_configure("bold", font=(input_text_font.actual()["family"], input_text_font.actual()["size"], "bold"))
input_text.tag_configure("italic", font=(input_text_font.actual()["family"], input_text_font.actual()["size"], "italic"))
input_text.tag_configure("underline", font=(input_text_font.actual()["family"], input_text_font.actual()["size"], "underline"))

# Botones y menús desplegables para aplicar estilos
style_frame = tk.Frame(app)
style_frame.pack(pady=5)

bold_button = tk.Button(style_frame, text="Negrita", command=lambda: apply_style("bold"))
bold_button.grid(row=0, column=0, padx=5)

italic_button = tk.Button(style_frame, text="Cursiva", command=lambda: apply_style("italic"))
italic_button.grid(row=0, column=1, padx=5)  # Corregido

underline_button = tk.Button(style_frame, text="Subrayado", command=lambda: apply_style("underline"))
underline_button.grid(row=0, column=2, padx=5)  # Corregido


align_label = tk.Label(style_frame, text="Alineación:")
align_label.grid(row=0, column=3, padx=5)

align_var = tk.StringVar()
align_var.set("left")
align_menu = ttk.OptionMenu(style_frame, align_var, "left", "left", "center", "right", "justify", command=lambda _: change_alignment())
align_menu.grid(row=0, column=4, padx=5)

# Botón para convertir texto ingresado
convert_button = tk.Button(app, text="Convertir texto", command=convert_text)
convert_button.pack(pady=5)

# Cuadro de texto para mostrar el resultado
output_label = tk.Label(app, text="Resultado:")
output_label.pack(pady=5)
output_text = tk.Text(app, wrap=tk.WORD, height=10)
output_text.pack(padx=10, pady=5)

app.mainloop()